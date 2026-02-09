"""
Generate ASPR_Photo_Repository_Requirements.docx with ASPR/HHS branding.
Run:  python scripts/generate-requirements-docx.py
"""

import os
from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

ROOT = Path(__file__).resolve().parent.parent
OUT = ROOT / "docs" / "ASPR_Photo_Repository_Requirements_v1.docx"

# ── ASPR / HHS brand colours ──────────────────────────────────────────
BLUE_DARK   = RGBColor(0x06, 0x2E, 0x61)
BLUE_PRIMARY = RGBColor(0x15, 0x51, 0x97)
GOLD        = RGBColor(0xAA, 0x64, 0x04)
RED         = RGBColor(0x99, 0x00, 0x00)
WHITE       = RGBColor(0xFF, 0xFF, 0xFF)
LIGHT_GRAY  = RGBColor(0xF2, 0xF2, 0xF2)

BLUE_DARK_HEX   = "062E61"
BLUE_PRIMARY_HEX = "155197"
GOLD_HEX        = "AA6404"
LIGHT_GRAY_HEX  = "F2F2F2"


# ── Helpers ───────────────────────────────────────────────────────────

def set_cell_shading(cell, hex_color):
    """Apply background shading to a table cell."""
    shading = parse_xml(
        f'<w:shd {nsdecls("w")} w:fill="{hex_color}" w:val="clear"/>'
    )
    cell._tc.get_or_add_tcPr().append(shading)


def set_cell_border(cell, **kwargs):
    """Set borders on a cell. kwargs: top, bottom, left, right with values like ('single','4','auto')."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = parse_xml(f'<w:tcBorders {nsdecls("w")}></w:tcBorders>')
    for edge, (style, sz, color) in kwargs.items():
        el = parse_xml(
            f'<w:{edge} {nsdecls("w")} w:val="{style}" w:sz="{sz}" w:color="{color}" w:space="0"/>'
        )
        tcBorders.append(el)
    tcPr.append(tcBorders)


def styled_table(doc, headers, rows, col_widths=None):
    """Create a branded table with dark-blue header row and alternating shading."""
    ncols = len(headers)
    table = doc.add_table(rows=1 + len(rows), cols=ncols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    # Header row
    hdr = table.rows[0]
    for i, text in enumerate(headers):
        cell = hdr.cells[i]
        set_cell_shading(cell, BLUE_DARK_HEX)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(text)
        run.bold = True
        run.font.color.rgb = WHITE
        run.font.size = Pt(9.5)
        run.font.name = "Calibri"
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after = Pt(3)

    # Data rows
    for ri, row_data in enumerate(rows):
        row = table.rows[ri + 1]
        for ci, text in enumerate(row_data):
            cell = row.cells[ci]
            if ri % 2 == 1:
                set_cell_shading(cell, LIGHT_GRAY_HEX)
            p = cell.paragraphs[0]
            run = p.add_run(str(text))
            run.font.size = Pt(9.5)
            run.font.name = "Calibri"
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)

    # Column widths
    if col_widths:
        total = sum(col_widths)
        table_width = Inches(6.5)
        for row in table.rows:
            for ci, cell in enumerate(row.cells):
                cell.width = int(table_width * col_widths[ci] / total)

    return table


def add_heading_styled(doc, text, level=1):
    """Add a heading with ASPR brand colour."""
    h = doc.add_heading(text, level=level)
    color_map = {1: BLUE_DARK, 2: BLUE_PRIMARY, 3: GOLD}
    color = color_map.get(level, BLUE_DARK)
    for run in h.runs:
        run.font.color.rgb = color
        run.font.name = "Calibri"
    return h


def add_para(doc, text, bold=False, italic=False, size=Pt(11), color=None, align=None, space_after=Pt(6)):
    """Add a paragraph with optional styling."""
    p = doc.add_paragraph()
    if align:
        p.alignment = align
    p.paragraph_format.space_after = space_after
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = size
    run.font.name = "Calibri"
    if color:
        run.font.color.rgb = color
    return p


def add_bullet(doc, text, level=0):
    """Add a bullet point."""
    p = doc.add_paragraph(style="List Bullet" if level == 0 else "List Bullet 2")
    p.paragraph_format.space_after = Pt(3)
    # Clear default run and add styled one
    p.clear()
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.font.name = "Calibri"
    return p


def add_page_break(doc):
    doc.add_page_break()


# ══════════════════════════════════════════════════════════════════════
#  BUILD DOCUMENT
# ══════════════════════════════════════════════════════════════════════

doc = Document()
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)
style.paragraph_format.space_after = Pt(6)

# Page margins
for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

# ── Header & Footer ──
section = doc.sections[0]

# Header
header = section.header
header.is_linked_to_previous = False
hp = header.paragraphs[0]
hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
run = hp.add_run("ASPR Photo Repository — Software Requirements Document")
run.italic = True
run.font.size = Pt(8)
run.font.color.rgb = BLUE_PRIMARY
run.font.name = "Calibri"

# Footer
footer = section.footer
footer.is_linked_to_previous = False
fp = footer.paragraphs[0]
fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = fp.add_run("HHS/ASPR — For Official Use Only")
run.font.size = Pt(8)
run.font.color.rgb = BLUE_DARK
run.font.name = "Calibri"


# ══════════════════════════════════════════════════════════════════════
#  COVER PAGE
# ══════════════════════════════════════════════════════════════════════

# Spacer
for _ in range(4):
    doc.add_paragraph()

# ASPR logo
logo_path = ROOT / "public" / "aspr-logo-blue.png"
if logo_path.exists():
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(logo_path), width=Inches(2.5))

doc.add_paragraph()

add_para(doc, "U.S. Department of Health and Human Services",
         size=Pt(12), color=BLUE_DARK, align=WD_ALIGN_PARAGRAPH.CENTER)
add_para(doc, "Administration for Strategic Preparedness and Response (ASPR)",
         size=Pt(11), color=BLUE_PRIMARY, align=WD_ALIGN_PARAGRAPH.CENTER)

for _ in range(2):
    doc.add_paragraph()

add_para(doc, "Software Requirements Document",
         bold=True, size=Pt(26), color=BLUE_DARK, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(8))
add_para(doc, "ASPR Photo Repository Application",
         size=Pt(18), color=BLUE_PRIMARY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(24))

for _ in range(2):
    doc.add_paragraph()

# Document info table
styled_table(doc,
    ["Property", "Value"],
    [
        ["Document Version", "1.0"],
        ["Date", "February 6, 2026"],
        ["Application Version", "0.1.0"],
        ["Project", "app-aspr-photos-lab"],
        ["Status", "Development"],
        ["Classification", "For Official Use Only"],
    ],
    col_widths=[35, 65],
)

for _ in range(3):
    doc.add_paragraph()

add_para(doc, "DRAFT — FOR REVIEW",
         bold=True, size=Pt(14), color=RED, align=WD_ALIGN_PARAGRAPH.CENTER)

add_page_break(doc)


# ══════════════════════════════════════════════════════════════════════
#  TABLE OF CONTENTS
# ══════════════════════════════════════════════════════════════════════

add_heading_styled(doc, "Table of Contents", level=1)

# TOC field — Word will populate when user presses "Update field"
p = doc.add_paragraph()
run = p.add_run()
fldChar1 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
run._r.append(fldChar1)

run2 = p.add_run()
instrText = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> TOC \\o "1-3" \\h \\z \\u </w:instrText>')
run2._r.append(instrText)

run3 = p.add_run()
fldChar2 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="separate"/>')
run3._r.append(fldChar2)

run4 = p.add_run("[Right-click and select 'Update Field' to generate Table of Contents]")
run4.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
run4.font.size = Pt(10)
run4.italic = True

run5 = p.add_run()
fldChar3 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
run5._r.append(fldChar3)

add_page_break(doc)


# ══════════════════════════════════════════════════════════════════════
#  1. INTRODUCTION
# ══════════════════════════════════════════════════════════════════════

add_heading_styled(doc, "1. Introduction", level=1)

add_heading_styled(doc, "1.1 Purpose", level=2)
add_para(doc,
    "The ASPR Photo Repository is a secure, web-based application that enables "
    "Administration for Strategic Preparedness and Response (ASPR) field teams to capture, upload, and "
    "manage disaster-related photographs during incident response operations."
)

add_heading_styled(doc, "1.2 Scope", level=2)
add_para(doc, "The application provides:")
add_bullet(doc, "PIN-based authentication for field teams")
add_bullet(doc, "Photo upload with metadata (GPS coordinates, incident ID, notes)")
add_bullet(doc, "Admin dashboard for session and PIN management")
add_bullet(doc, "Secure storage via Azure cloud services")
add_bullet(doc, "Government-compliant (ASPR/HHS) branding and security controls")

add_heading_styled(doc, "1.3 Intended Users", level=2)
styled_table(doc,
    ["Role", "Description"],
    [
        ["Field Team Member", "ASPR responders who capture and upload photos from the field"],
        ["Admin", "Operations staff who create PINs, manage sessions, and monitor activity"],
    ],
    col_widths=[30, 70],
)

add_heading_styled(doc, "1.4 Technology Stack", level=2)
styled_table(doc,
    ["Component", "Technology"],
    [
        ["Framework", "Next.js 16.1.6 (React 19, TypeScript)"],
        ["Styling", "Tailwind CSS 4, shadcn/ui, Radix UI"],
        ["Database", "Azure SQL Server (mssql)"],
        ["Storage", "Azure Blob Storage"],
        ["Authentication", "JWT (jsonwebtoken), PIN-based"],
        ["Image Processing", "Sharp"],
        ["Deployment", "Azure App Service via GitHub Actions"],
        ["Runtime", "Node.js 22.x"],
    ],
    col_widths=[30, 70],
)


# ══════════════════════════════════════════════════════════════════════
#  2. FUNCTIONAL REQUIREMENTS
# ══════════════════════════════════════════════════════════════════════

add_page_break(doc)
add_heading_styled(doc, "2. Functional Requirements", level=1)

# 2.1 Auth
add_heading_styled(doc, "2.1 Authentication & Authorization", level=2)

add_heading_styled(doc, "FR-2.1.1 PIN Login (Field Teams)", level=3)
for req in [
    "The system SHALL present a login page at the root URL (/) with a 6-digit numeric PIN input.",
    "The system SHALL validate PINs against the upload_sessions database table.",
    "The system SHALL only accept PINs that are active (is_active = 1) and not expired (expires_at > current time).",
    "The system SHALL generate a JWT token (24-hour expiration) upon successful PIN validation.",
    "The system SHALL store the JWT token, session ID, and team name in the browser\u2019s sessionStorage.",
    "The system SHALL redirect authenticated users to the /upload page.",
    "The system SHALL display remaining login attempts on failure.",
]:
    add_bullet(doc, req)

add_heading_styled(doc, "FR-2.1.2 Admin Authentication", level=3)
for req in [
    "The system SHALL provide an admin dashboard at /admin.",
    "The system SHALL require Microsoft Entra ID SSO authentication for all administrative API calls.",
    "The system SHALL support admin login through the web dashboard UI.",
]:
    add_bullet(doc, req)

add_heading_styled(doc, "FR-2.1.3 Session Management", level=3)
for req in [
    "JWT tokens SHALL expire after 24 hours.",
    "Session data SHALL be stored in sessionStorage (cleared on tab close).",
    "PINs SHALL expire 7 days after creation.",
    "Users SHALL be able to log out, clearing all session data.",
]:
    add_bullet(doc, req)

# 2.2 Photo Upload
add_heading_styled(doc, "2.2 Photo Upload", level=2)

add_heading_styled(doc, "FR-2.2.1 File Upload", level=3)
for req in [
    "The system SHALL accept photo uploads at POST /api/photos/upload.",
    "The system SHALL require a valid JWT token in the Authorization: Bearer header.",
    "The system SHALL accept JPEG, PNG, and WebP image formats only.",
    "The system SHALL enforce a maximum file size of 50 MB per upload.",
    "The system SHALL validate filenames to allow only alphanumeric characters, spaces, hyphens, dots, and underscores.",
]:
    add_bullet(doc, req)

add_heading_styled(doc, "FR-2.2.2 Photo Metadata", level=3)
add_para(doc, "The system SHALL accept the following optional metadata with each upload:")
for req in [
    "Incident ID \u2014 format: XX-YYYY-000 (2 uppercase letters, dash, 4 digits, dash, 3 digits)",
    "GPS Coordinates \u2014 latitude (-90 to 90) and longitude (-180 to 180)",
    "Location Name \u2014 formatted coordinate string",
    "Notes \u2014 free text, maximum 1,000 characters",
]:
    add_bullet(doc, req)

add_heading_styled(doc, "FR-2.2.3 Image Processing", level=3)
for req in [
    "The system SHALL extract image metadata (width, height, format) using Sharp.",
    "The system SHALL generate a thumbnail for each uploaded photo:",
]:
    add_bullet(doc, req)
for req in [
    "Maximum dimensions: 400 \u00d7 300 pixels",
    "Format: WebP, Quality: 80",
    "Aspect ratio preserved, no upscaling",
]:
    add_bullet(doc, req, level=1)

add_heading_styled(doc, "FR-2.2.4 Storage", level=3)
for req in [
    "The system SHALL upload the original photo to Azure Blob Storage at path aspr-photos/{photoId}/original.",
    "The system SHALL upload the generated thumbnail at path aspr-photos/{photoId}/thumbnail.",
    "The system SHALL store photo metadata (file name, blob URL, file size, dimensions, MIME type, coordinates, notes, incident ID, timestamp) in the photos database table.",
    "The system SHALL return a success response with photoId and file size.",
]:
    add_bullet(doc, req)

add_heading_styled(doc, "FR-2.2.5 Upload User Interface", level=3)
for req in [
    "The upload page SHALL display the authenticated team name.",
    "The upload page SHALL provide a camera/file selection button.",
    "The upload page SHALL provide optional fields for incident ID, location, and notes.",
    "The upload page SHALL support browser geolocation for GPS coordinates.",
    "The upload page SHALL show upload progress and success/error notifications.",
    "The upload page SHALL prompt the user to capture another photo after a successful upload.",
]:
    add_bullet(doc, req)

# 2.3 Admin Dashboard
add_heading_styled(doc, "2.3 Admin Dashboard", level=2)

add_heading_styled(doc, "FR-2.3.1 PIN Creation", level=3)
for req in [
    "Admins SHALL be able to create new 6-digit PINs via the web dashboard, CLI tool, or direct API call.",
    "The system SHALL generate random 6-digit PINs (range: 100000\u2013999999).",
    "Admins SHALL be able to assign a team name to each PIN (auto-generated if left blank).",
    "Created PINs SHALL have a 7-day expiration.",
]:
    add_bullet(doc, req)

add_heading_styled(doc, "FR-2.3.2 PIN Management", level=3)
for req in [
    "The admin dashboard SHALL display all created PINs with their team names and expiration dates.",
    "The admin dashboard SHALL provide copy-to-clipboard functionality for PINs.",
]:
    add_bullet(doc, req)

add_heading_styled(doc, "FR-2.3.3 Admin API", level=3)
for req in [
    "POST /api/auth/create-session SHALL create a new PIN and return { id, pin, team_name }.",
    "The endpoint SHALL require Entra ID session authentication.",
    "The endpoint SHALL validate team names (max 255 characters, alphanumeric with spaces/hyphens/underscores).",
]:
    add_bullet(doc, req)


# ══════════════════════════════════════════════════════════════════════
#  3. NON-FUNCTIONAL REQUIREMENTS
# ══════════════════════════════════════════════════════════════════════

add_page_break(doc)
add_heading_styled(doc, "3. Non-Functional Requirements", level=1)

# 3.1 Security
add_heading_styled(doc, "3.1 Security", level=2)

add_heading_styled(doc, "NFR-3.1.1 Rate Limiting", level=3)
styled_table(doc,
    ["Endpoint", "Max Attempts", "Window", "Lockout"],
    [
        ["PIN Validation (/api/auth/validate-pin)", "5", "60 seconds", "15 minutes"],
        ["Admin Auth (failed attempts)", "3", "60 seconds", "30 minutes"],
        ["PIN Creation (/api/auth/create-session)", "20", "60 seconds", "None"],
        ["Photo Upload (/api/photos/upload)", "50", "1 hour", "None"],
    ],
    col_widths=[40, 20, 20, 20],
)
add_bullet(doc, "Rate limiting SHALL be enforced per IP address.")
add_bullet(doc, "Rate limit state SHALL be stored in-memory (with recommendation to migrate to Redis for production scale).")

add_heading_styled(doc, "NFR-3.1.2 Input Validation", level=3)
for req in [
    "All user inputs SHALL be validated server-side before processing.",
    "PINs SHALL match the pattern: exactly 6 numeric digits.",
    "Incident IDs SHALL match the pattern: XX-YYYY-000.",
    "All database queries SHALL use parameterized inputs to prevent SQL injection.",
]:
    add_bullet(doc, req)

add_heading_styled(doc, "NFR-3.1.3 Security Headers", level=3)
add_para(doc, "The application SHALL set the following HTTP headers on all responses:")
styled_table(doc,
    ["Header", "Value"],
    [
        ["X-Content-Type-Options", "nosniff"],
        ["X-Frame-Options", "DENY"],
        ["X-XSS-Protection", "1; mode=block"],
        ["Strict-Transport-Security", "max-age=31536000; includeSubDomains; preload"],
        ["Referrer-Policy", "strict-origin-when-cross-origin"],
        ["Permissions-Policy", "camera=(), microphone=(), geolocation=(self), payment=()"],
        ["Content-Security-Policy", "Restricts sources to \u2018self\u2019 with necessary exceptions"],
    ],
    col_widths=[35, 65],
)
add_bullet(doc, "The X-Powered-By header SHALL be suppressed.")

add_heading_styled(doc, "NFR-3.1.4 Audit Logging", level=3)
add_para(doc, "The system SHALL log the following security events with timestamp, IP address, and relevant details:")
styled_table(doc,
    ["Event", "Description"],
    [
        ["AUTH_SUCCESS", "Successful PIN validation"],
        ["AUTH_FAILURE", "Failed PIN validation"],
        ["PIN_CREATED", "New PIN created (last 2 digits only logged)"],
        ["UPLOAD_SUCCESS", "Successful photo upload"],
        ["UPLOAD_FAILURE", "Failed photo upload"],
        ["RATE_LIMIT_EXCEEDED", "Rate limit triggered"],
    ],
    col_widths=[35, 65],
)

add_heading_styled(doc, "NFR-3.1.5 Error Handling", level=3)
for req in [
    "The system SHALL NOT expose internal error details to clients.",
    "The system SHALL return generic, safe error messages (400, 401, 403, 404, 429, 500).",
    "The system SHALL log full error details server-side.",
]:
    add_bullet(doc, req)

add_heading_styled(doc, "NFR-3.1.6 OWASP Compliance", level=3)
add_para(doc, "The application SHALL address all OWASP Top 10 (2021) categories:")
styled_table(doc,
    ["#", "Category", "Implementation"],
    [
        ["1", "Broken Access Control", "JWT verification, Entra ID session validation"],
        ["2", "Cryptographic Failures", "HTTPS enforcement, JWT HS256"],
        ["3", "Injection", "Parameterized SQL, input validation"],
        ["4", "Insecure Design", "Rate limiting, session expiry"],
        ["5", "Security Misconfiguration", "Security headers, no debug output"],
        ["6", "Vulnerable Components", "Dependency auditing"],
        ["7", "Authentication Failures", "JWT expiration, session management"],
        ["8", "Software Integrity", "Version control, CI/CD pipeline"],
        ["9", "Logging & Monitoring", "Audit logging"],
        ["10", "SSRF", "No user-controlled URL fetching"],
    ],
    col_widths=[8, 32, 60],
)

# 3.2 Performance
add_heading_styled(doc, "3.2 Performance", level=2)
for req in [
    "Thumbnail generation SHALL complete server-side using Sharp without blocking the upload response unnecessarily.",
    "The system SHALL support files up to 50 MB.",
    "The application SHALL use Next.js standalone output mode for minimal deployment footprint.",
    "Source maps SHALL be disabled in production.",
    "Response compression SHALL be enabled.",
]:
    add_bullet(doc, req)

# 3.3 Reliability
add_heading_styled(doc, "3.3 Reliability", level=2)
for req in [
    "In development mode, the application SHALL fall back to an in-memory mock database if Azure SQL Server is unavailable.",
    "The system SHALL automatically create the aspr-photos blob container if it does not exist.",
]:
    add_bullet(doc, req)

# 3.4 Usability
add_heading_styled(doc, "3.4 Usability", level=2)
for req in [
    "The application SHALL use ASPR/HHS government branding (Primary Blue #155197, Dark Blue #062e61, Gold #AA6404, Red #990000).",
    "The UI SHALL use the shadcn/ui component library with Tailwind CSS.",
    "The application SHALL be mobile-responsive (mobile-first design).",
    "The application SHALL use semantic HTML with proper focus states and ARIA labels.",
    "Typography SHALL use Geist Sans (primary) and Geist Mono (monospace).",
    "Icons SHALL use the lucide-react library.",
]:
    add_bullet(doc, req)

# 3.5 Deployment
add_heading_styled(doc, "3.5 Deployment & Infrastructure", level=2)
styled_table(doc,
    ["Azure Service", "Purpose"],
    [
        ["Azure App Service", "Application hosting (Node.js 22.x, Linux)"],
        ["Azure SQL Database", "Session and photo metadata storage"],
        ["Azure Blob Storage", "Photo and thumbnail file storage"],
        ["Azure Key Vault", "Secrets management (recommended)"],
    ],
    col_widths=[35, 65],
)
for req in [
    "The application SHALL deploy automatically via GitHub Actions on push to the main branch.",
    "Azure authentication SHALL use OpenID Connect (OIDC) with federated identity.",
]:
    add_bullet(doc, req)

add_heading_styled(doc, "NFR-3.5.1 Environment Configuration", level=3)
styled_table(doc,
    ["Variable", "Description", "Required"],
    [
        ["JWT_SECRET", "Secret key for JWT signing", "Yes"],
        ["AUTH_MICROSOFT_ENTRA_ID_ID", "Entra ID app client ID", "Yes"],
        ["SQL_SERVER", "Azure SQL Database server hostname", "Yes"],
        ["SQL_DATABASE", "Database name", "Yes"],
        ["SQL_USERNAME", "Database username (dev only)", "Dev"],
        ["SQL_PASSWORD", "Database password (dev only)", "Dev"],
        ["AZURE_STORAGE_CONNECTION_STRING", "Blob Storage connection string", "Yes"],
        ["NODE_ENV", "Runtime environment", "No"],
        ["NEXT_PUBLIC_API_URL", "Public API base URL", "No"],
    ],
    col_widths=[40, 45, 15],
)


# ══════════════════════════════════════════════════════════════════════
#  4. DATA MODEL
# ══════════════════════════════════════════════════════════════════════

add_page_break(doc)
add_heading_styled(doc, "4. Data Model", level=1)

add_heading_styled(doc, "4.1 upload_sessions", level=2)
styled_table(doc,
    ["Column", "Type", "Constraints", "Description"],
    [
        ["id", "NVARCHAR(36)", "PRIMARY KEY", "UUID"],
        ["pin", "NVARCHAR(6)", "NOT NULL", "6-digit PIN"],
        ["team_name", "NVARCHAR(255)", "NOT NULL", "Team identifier"],
        ["is_active", "BIT", "DEFAULT 1", "Active flag"],
        ["created_at", "DATETIME", "DEFAULT GETDATE()", "Creation timestamp"],
        ["expires_at", "DATETIME", "NOT NULL", "Expiration (7 days from creation)"],
    ],
    col_widths=[20, 25, 25, 30],
)

add_heading_styled(doc, "4.2 photos", level=2)
styled_table(doc,
    ["Column", "Type", "Constraints", "Description"],
    [
        ["id", "NVARCHAR(36)", "PRIMARY KEY", "UUID"],
        ["session_id", "NVARCHAR(36)", "FK \u2192 upload_sessions.id", "Owning session"],
        ["file_name", "NVARCHAR(255)", "NOT NULL", "Original filename"],
        ["blob_url", "NVARCHAR(MAX)", "NOT NULL", "Azure Blob Storage URL"],
        ["file_size", "BIGINT", "NOT NULL", "Size in bytes"],
        ["width", "INT", "NULL", "Image width (px)"],
        ["height", "INT", "NULL", "Image height (px)"],
        ["mime_type", "NVARCHAR(50)", "NULL", "MIME type"],
        ["latitude", "FLOAT", "NULL", "GPS latitude"],
        ["longitude", "FLOAT", "NULL", "GPS longitude"],
        ["location_name", "NVARCHAR(255)", "NULL", "Formatted location"],
        ["notes", "NVARCHAR(1000)", "NULL", "User notes"],
        ["incident_id", "NVARCHAR(50)", "NULL", "Incident identifier"],
        ["timestamp", "DATETIME", "DEFAULT GETDATE()", "Upload timestamp"],
    ],
    col_widths=[20, 25, 25, 30],
)

add_heading_styled(doc, "4.3 Relationships", level=2)
add_bullet(doc, "One upload_session \u2192 Many photos (via session_id foreign key)")
add_bullet(doc, "A single PIN/session can be shared across multiple team members.")


# ══════════════════════════════════════════════════════════════════════
#  5. API SPECIFICATION
# ══════════════════════════════════════════════════════════════════════

add_page_break(doc)
add_heading_styled(doc, "5. API Specification", level=1)

add_heading_styled(doc, "5.1 POST /api/auth/validate-pin", level=2)
add_para(doc, "Authenticate a field team member with a PIN.")
styled_table(doc,
    ["Property", "Details"],
    [
        ["Authentication", "None"],
        ["Request Body", '{ "pin": "123456" }'],
        ["Success Response (200)", '{ "sessionId": "uuid", "teamName": "string", "token": "jwt" }'],
        ["Error 400", "Invalid PIN format"],
        ["Error 401", "Invalid or expired PIN"],
        ["Error 429", "Rate limited"],
    ],
    col_widths=[30, 70],
)

add_heading_styled(doc, "5.2 POST /api/auth/create-session", level=2)
add_para(doc, "Create a new PIN (admin only).")
styled_table(doc,
    ["Property", "Details"],
    [
        ["Authentication", "Entra ID session (required)"],
        ["Request Body", '{ "teamName": "Team A" } (optional)'],
        ["Success Response (200)", '{ "id": "uuid", "pin": "654321", "team_name": "Team A" }'],
        ["Error 401", "Unauthenticated"],
        ["Error 429", "Rate limited"],
    ],
    col_widths=[30, 70],
)

add_heading_styled(doc, "5.3 POST /api/photos/upload", level=2)
add_para(doc, "Upload a photo with metadata.")
styled_table(doc,
    ["Property", "Details"],
    [
        ["Authentication", "Authorization: Bearer {JWT}"],
        ["Content-Type", "multipart/form-data"],
        ["Required Fields", "photo (JPEG/PNG/WebP, max 50 MB)"],
        ["Optional Fields", "notes, incidentId, latitude, longitude, locationName"],
        ["Success Response (200)", '{ "success": true, "photoId": "uuid", "size": "X.XX MB" }'],
        ["Error 400", "Invalid file or metadata"],
        ["Error 401", "Unauthorized"],
        ["Error 429", "Rate limited"],
    ],
    col_widths=[30, 70],
)


# ══════════════════════════════════════════════════════════════════════
#  6. SYSTEM LIMITS
# ══════════════════════════════════════════════════════════════════════

add_page_break(doc)
add_heading_styled(doc, "6. System Limits", level=1)

styled_table(doc,
    ["Parameter", "Value"],
    [
        ["PIN length", "6 digits"],
        ["PIN expiration", "7 days"],
        ["JWT token expiration", "24 hours"],
        ["Max upload file size", "50 MB"],
        ["Supported image types", "JPEG, PNG, WebP"],
        ["Thumbnail max dimensions", "400 \u00d7 300 px"],
        ["Thumbnail format", "WebP (quality 80)"],
        ["Max notes length", "1,000 characters"],
        ["Max team name length", "255 characters"],
        ["Latitude range", "\u221290 to 90"],
        ["Longitude range", "\u2212180 to 180"],
    ],
    col_widths=[40, 60],
)


# ══════════════════════════════════════════════════════════════════════
#  7. PAGES & NAVIGATION
# ══════════════════════════════════════════════════════════════════════

add_heading_styled(doc, "7. Pages & Navigation", level=1)

styled_table(doc,
    ["Route", "Page", "Access", "Description"],
    [
        ["/", "Login", "Public", "PIN entry form, redirects to /upload on success"],
        ["/upload", "Photo Upload", "Authenticated (JWT)", "Photo capture/upload with metadata fields"],
        ["/admin", "Admin Dashboard", "Entra ID SSO", "PIN creation and management"],
    ],
    col_widths=[15, 20, 25, 40],
)


# ══════════════════════════════════════════════════════════════════════
#  8. FUTURE CONSIDERATIONS
# ══════════════════════════════════════════════════════════════════════

add_heading_styled(doc, "8. Future Considerations", level=1)

add_para(doc,
    "The following features are not currently implemented but are supported "
    "by the existing data model and architecture:"
)
for item in [
    "Photo Gallery / Viewer \u2014 browse uploaded photos by session, incident, or date",
    "Geospatial Map View \u2014 plot photos on a map using stored GPS coordinates",
    "Photo Search & Filtering \u2014 search by incident ID, date range, or team",
    "PIN Revocation \u2014 admin ability to deactivate PINs before expiration",
    "Redis Rate Limiting \u2014 distributed rate limiting for multi-instance deployments",
    "OAuth 2.0 / Entra ID Integration \u2014 enterprise SSO for admin access",
    "Dark Mode \u2014 theme toggle using the existing design system",
    "Structured Audit Log Table \u2014 persistent database-backed audit trail",
    "Azure Application Insights \u2014 production monitoring and alerting",
]:
    add_bullet(doc, item)


# ══════════════════════════════════════════════════════════════════════
#  9. PROJECT STRUCTURE
# ══════════════════════════════════════════════════════════════════════

add_page_break(doc)
add_heading_styled(doc, "9. Project Structure", level=1)

structure = [
    ("app-aspr-photos-lab/", ""),
    ("  app/", "Next.js app directory"),
    ("    page.tsx", "Login page (/)"),
    ("    layout.tsx", "Root layout"),
    ("    globals.css", "Global styles"),
    ("    upload/page.tsx", "Photo upload page"),
    ("    admin/page.tsx", "Admin dashboard"),
    ("    api/auth/create-session/", "PIN creation API"),
    ("    api/auth/validate-pin/", "PIN validation API"),
    ("    api/photos/upload/", "Photo upload API"),
    ("  components/ui/", "shadcn/ui components"),
    ("    button.tsx, input.tsx, card.tsx, alert.tsx, textarea.tsx", ""),
    ("  lib/", "Shared libraries"),
    ("    auth.ts", "JWT sign/verify"),
    ("    db.ts", "Database connection & queries"),
    ("    rateLimit.ts", "Rate limiting"),
    ("    security.ts", "Validation & audit logging"),
    ("    utils.ts", "Utility functions"),
    ("  docs/", "Documentation"),
    ("  scripts/", "Build & generation scripts"),
    ("  .github/workflows/", "CI/CD pipeline"),
    ("  SECURITY.md", "Security documentation"),
    ("  DESIGN_SYSTEM.md", "UI component documentation"),
]

styled_table(doc,
    ["Path", "Description"],
    [[p, d] for p, d in structure],
    col_widths=[50, 50],
)


# ══════════════════════════════════════════════════════════════════════
#  10. DOCUMENT APPROVAL
# ══════════════════════════════════════════════════════════════════════

add_page_break(doc)
add_heading_styled(doc, "10. Document Approval", level=1)

add_para(doc, "This document requires review and approval from the following stakeholders:")

styled_table(doc,
    ["Role", "Name", "Signature", "Date"],
    [
        ["Project Sponsor", "", "", ""],
        ["Technical Lead", "", "", ""],
        ["Security Officer", "", "", ""],
        ["Operations Lead", "", "", ""],
    ],
    col_widths=[25, 25, 30, 20],
)

doc.add_paragraph()

add_heading_styled(doc, "Revision History", level=2)
styled_table(doc,
    ["Version", "Date", "Author", "Changes"],
    [
        ["1.0", "February 6, 2026", "", "Initial requirements document"],
    ],
    col_widths=[15, 20, 25, 40],
)


# ══════════════════════════════════════════════════════════════════════
#  SAVE
# ══════════════════════════════════════════════════════════════════════

doc.save(str(OUT))
size_kb = OUT.stat().st_size / 1024
print(f"\nDocument generated: {OUT}")
print(f"Size: {size_kb:.1f} KB")
